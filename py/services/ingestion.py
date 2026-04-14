from typing import Union, BinaryIO, IO, Optional, List
from datetime import datetime, date
import fitz
from docx import Document
from custom.structer import DocumentMetaData
import logging


class IngestionService:
    def __init__(self):
        pass

    def _ingest_pdf(self, file: Union[str, BinaryIO, IO]) -> Optional[DocumentMetaData]:
        doc = None
        try:
            doc = fitz.open(file)

            text_parts: List[str] = []
            for page in doc:
                page_text = page.get_text("text")  # simple + reliable
                if page_text:
                    text_parts.append(page_text.strip())

            metadata = doc.metadata or {}

            # Try parsing date safely
            raw_date = metadata.get("creationDate")
            parsed_date = self._parse_pdf_date(raw_date)

            return DocumentMetaData(
                name=self._get_filename(file),
                title=metadata.get("title"),
                date_conducted=parsed_date,
                page_nums=len(doc),
                text="\n\n".join(text_parts)
            )

        except Exception as e:
            print(f"[PDF ERROR] {e}")
            return None

        finally:
            if doc:
                doc.close()

    def _ingest_docx(self, file: Union[str, BinaryIO, IO]) -> Optional[DocumentMetaData]:
        try:
            doc = Document(file)

            text_parts: List[str] = []

            for para in doc.paragraphs:
                clean = para.text.strip()
                if clean:
                    text_parts.append(clean)

            core = doc.core_properties

            return DocumentMetaData(
                name=self._get_filename(file),
                title=core.title,
                date_conducted=core.created.date() if core.created else None,
                page_nums=len(doc.paragraphs),  # approximation
                text="\n\n".join(text_parts)
            )

        except Exception as e:
            print(f"[DOCX ERROR] {e}")
            return None

    def ingest(self, file: Union[str, BinaryIO, IO]) -> Optional[DocumentMetaData]:
        if isinstance(file, str):
            ext = file.lower().split('.')[-1]
        else:
            ext = ""

        if ext == "pdf":
            return self._ingest_pdf(file)
        elif ext == "docx":
            return self._ingest_docx(file)
        else:
            raise ValueError("Unsupported file type")


    def _get_filename(self, file: Union[str, BinaryIO, IO]) -> str:
        if isinstance(file, str):
            return file.split("/")[-1]
        return "uploaded_file"

    def _parse_pdf_date(self, raw: Optional[str]) -> Optional[date]:
        if not raw:
            return None

        try:
            raw = raw.replace("D:", "")
            dt = datetime.strptime(raw[:14], "%Y%m%d%H%M%S")
            return dt.date()
        except Exception:
            return None